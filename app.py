from flask import Flask, render_template, request, make_response, abort, send_file
import psycopg2
import psycopg2.extras
import json
from decimal import Decimal
from docx import Document
from io import BytesIO
import os

DB_USER_NAME = os.environ['DB_USER_NAME']
DB_PASSWORD = os.environ['DB_PASSWORD']
DB_DATABASE = os.environ['DB_DATABASE']
DB_HOST = os.environ['DB_HOST']

#Databse connection string
dbcon = psycopg2.connect(
        user=DB_USER_NAME,
        password=DB_PASSWORD,
        database=DB_DATABASE,
        host=DB_HOST)

#Create flask app object
app = Flask(__name__)

def default(obj):
    #Helper function for Decimal to JSON conversion
    if isinstance(obj, Decimal):
        return str(obj)
    raise TypeError("Object of type '%s' is not JSON serializable" % type(obj).__name__)

#Web app routes
@app.route("/")
def form():
    #Return Car Search
    return render_template('Car_search.html')

@app.route("/car/<int:id>")
def car(id):
    #Return details of a particular car

    #Create a databse cursor
    cur = dbcon.cursor()

    #Excecute a sql query to get the car details using  cursor
    cur.execute("SELECT * FROM cars WHERE id=%s", (id,))
    #Fetch the first returned tuple into car
    car = cur.fetchone()

    #Excecute a sql query to get the company details using  cursor
    cur.execute("SELECT * FROM company WHERE name=%s", (car[2],))
    #Fetch the first returned tuple into company
    company = cur.fetchone()

    #close the database cursor
    cur.close()

    #return the Car.html tmeplate with substituted values from company and car
    return render_template('Car.html', company=company, car=car)

@app.route("/results", methods=['GET', 'POST'])
def results():
    #Return details of Car Search Results
    if request.method == 'POST':
        #Check if HTTP request method is POST

        #Get fields from submitted form using request.form.get()
        company = request.form.get('company')

        price_starting = float(request.form.get('price_starting'))
        price_topend = float(request.form.get('price_topend'))

        manual = request.form.get('manual')
        if manual:
            manual = True

        automatic = request.form.get('automatic')
        if automatic:
            automatic = True

        petrol = request.form.get('petrol')
        if petrol:
            petrol = True

        diesel = request.form.get('diesel')
        if diesel:
            diesel = True

        cng = request.form.get('cng')
        if cng:
            cng = True

        electric = request.form.get('electric')
        if electric:
            cng = electric

        seating_capacity = int(request.form.get('seating_capacity'))

        car_sql = """
            SELECT id, name, company, image, price_starting, price_topend, mileage_l, mileage_u, manual, automatic, petrol, diesel, cng, electric, seating_capacity
            FROM cars
            WHERE company = %s
                AND price_starting > %s 
                AND price_topend < %s
                AND manual = COALESCE(%s, manual)
                AND automatic = COALESCE(%s, automatic)
                AND petrol = COALESCE(%s, petrol)
                AND diesel = COALESCE(%s, diesel)
                AND cng = COALESCE(%s, cng)
                AND electric = COALESCE(%s, electric)
                AND seating_capacity = %s;
                """
        car_params = (company, price_starting, price_topend, manual, automatic, petrol, diesel, cng, electric, seating_capacity)
        cur = dbcon.cursor()
        #Get result of cars that match the search criteria
        cur.execute(car_sql, car_params)
        cars = cur.fetchall()

        cur.execute("SELECT * FROM company WHERE name=%s", (company,))   
        company = cur.fetchone()

        return render_template("Home.html", company=company,cars=cars)

@app.route("/download/<int:id>")
def download_doc(id):
    #Return a docx file containing details of a particular car
    cur = dbcon.cursor()
    cur.execute('SELECT * FROM cars WHERE id=%s', (id,))
    car = cur.fetchone()
    cur.close()

    #Check if car with id exists
    if car is None:
        return '', 404

    #Create the docx file with the required car details 
    document = Document()
    head =document.add_heading(car[1],0)
    head.alignment = 1
    document.add_paragraph(str(car[4]))
    document.add_paragraph('Price range: '+str(car[5])+' lakh to '+str(car[6])+' lakh')
    document.add_paragraph('Mileage: '+str(car[7])+' kmpl '+'-'+str(car[8])+' kmpl')
    document.add_paragraph('Seating capacity: '+str(car[15]))
    f = BytesIO()
    document.save(f)
    f.seek(0)

    #Return the docx document
    return send_file(
        f,
        as_attachment=True,
        attachment_filename='car_doc.docx'
    )

# API routes
# localhost:5000/api/
# 
# Endpoints:
# GET /api/cars
# POST /api/cars
# GET /api/cars/:id
# PUT /api/cars:id
# DELETE /api/cars:id

@app.errorhandler(400)
def not_found(error):
    #Return JSON response for 400 errors
    r = make_response(json.dumps({ 'error': 'Bad request' }), 400)
    r.mimetype = 'application/json'
    return r

@app.errorhandler(404)
def not_found(error):
    #Return JSON response for 404 errors
    r = make_response(json.dumps({'error': 'Not found'}), 404)
    r.mimetype = 'application/json'
    return r

@app.route('/api/cars')
def get_cars():
    #GET /api/cars

    #psycopg2.extras.RealDictCursor returns rows as a Python Dictionary
    cur = dbcon.cursor(cursor_factory = psycopg2.extras.RealDictCursor)
    cur.execute('select * from cars')
    cars = cur.fetchall()
    cur.close()

    #Create a response with json body
    r = make_response(json.dumps({'cars': cars}, default=default))
    r.mimetype = 'application/json'
    return r

@app.route('/api/cars', methods = ['POST'])
def create_car():
    #POST /api/cars

    #Check if json string not present in body
    if not request.json:
        abort(400)
    
    #Check if all required fields are present in json
    for key in ['name' , 'company' , 'image' , 'summary' , 'price_starting' , 'price_topend' , 'mileage_l' , 'mileage_u' , 'manual' , 'automatic' , 'petrol' , 'diesel' , 'cng' , 'electric' , 'seating']:
        if key not in request.json:
            abort(400)

    car = request.json

    cur = dbcon.cursor(cursor_factory = psycopg2.extras.RealDictCursor)

    SQL = '''INSERT INTO cars(name, company, image, summary, price_starting, price_topend, mileage_l, mileage_u, manual, automatic, petrol, diesel, cng, electric, seating_capacity)
            VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING id'''
    
    params = (car['name'], car['company'], car['image'], car['summary'], car['price_starting'], car['price_topend'], car['mileage_l'], car['mileage_u'], car['manual'], car['automatic'], car['petrol'], car['diesel'], car['cng'], car['electric'], car['seating'])

    cur.execute(SQL, params)

    id = cur.fetchone()['id']
    cur.execute('SELECT * FROM cars WHERE id=%s', (id,))
    car = cur.fetchone()
    cur.close()

    #commit db transaction
    dbcon.commit()
    
    #create a json response with a status code 201
    r = make_response(json.dumps(car, default=default), 201)
    r.mimetype = 'application/json'
    return r

@app.route('/api/cars/<int:id>')
def get_car(id):
    #GET /api/cars/:id
    cur = dbcon.cursor(cursor_factory = psycopg2.extras.RealDictCursor)
    cur.execute('select * from cars where id=%s', (id,))
    car = cur.fetchone()
    cur.close()

    #Check if car with given id doesnt exist
    if car is None:
        abort(404)

    r = make_response(json.dumps(car, default=default))
    r.mimetype = 'application/json'
    return r

@app.route('/api/cars/<int:id>', methods = ['DELETE'])
def delete_car(id):
    #DELETE /api/cars/:id
    cur = dbcon.cursor(cursor_factory = psycopg2.extras.RealDictCursor)
    cur.execute('select * from cars where id=%s', (id,))
    car = cur.fetchone()

    #Check if car with given id doesnt exist
    if car is None:
        abort(404)

    cur.execute('delete from cars where id=%s', (id,))
    cur.close()

    #Return a successful response with empty body 
    return '', 204

@app.route('/api/cars/<int:id>', methods = ['PUT'])
def update_car(id):
    #PUT /api/cars/:id

    #Check if json string not present in body
    if not request.json:
        abort(400)
    
    #Check if all required fields are present in json
    for key in ['name' , 'company' , 'image' , 'summary' , 'price_starting' , 'price_topend' , 'mileage_l' , 'mileage_u' , 'manual' , 'automatic' , 'petrol' , 'diesel' , 'cng' , 'electric' , 'seating']:
        if key not in request.json:
            abort(400)

    cur = dbcon.cursor(cursor_factory = psycopg2.extras.RealDictCursor)
    cur.execute('select * from cars where id=%s', (id,))
    car = cur.fetchone()

    #check if car with given doesnt exist
    if car is None:
        abort(404)

    car = request.json

    SQL = '''
        UPDATE cars 
            SET name=%s,
                company=%s, 
                image=%s, 
                summary=%s, 
                price_starting=%s, 
                price_topend=%s, 
                mileage_l=%s, 
                mileage_u=%s, 
                manual=%s, 
                automatic=%s, 
                petrol=%s, 
                diesel=%s,
                cng=%s, 
                electric=%s, 
                seating_capacity=%s
            WHERE id=%s
    '''
    params = (car['name'], car['company'], car['image'], car['summary'], car['price_starting'], car['price_topend'], car['mileage_l'], car['mileage_u'], car['manual'], car['automatic'], car['petrol'], car['diesel'], car['cng'], car['electric'], car['seating'], id)

    #update car with given with new values
    cur.execute(SQL, params)
    cur.close()

    #Return a successful response with empty body
    return '', 204

if __name__ == "__main__":
    #Run the app
    app.run()

    #Close database connection
    dbcon.close()
